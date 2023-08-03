from collections import namedtuple
import altair as alt
import math
import pandas as pd
import seaborn as sns
import streamlit as st
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from io import BytesIO
from streamlit.components.v1 import html
from docx import Document
from docx.shared import Inches
import base64

def generate_report(df):
    # Créer un document Word
    doc = Document()
    doc.add_heading("Analysis report", level=1)

    # Ajouter le nombre d'enregistrements au rapport
    num_records = len(df)
    doc.add_paragraph(f"Number of records : {num_records}")

    # Ajouter le tableau issu de df.describe() au rapport
    doc.add_heading("General infos:", level=2)
    df_describe_table = df.describe().reset_index()
    df_describe_table.columns = [""] + list(df_describe_table.columns[1:])
    table = doc.add_table(df_describe_table.shape[0]+1, df_describe_table.shape[1])
    for i in range(df_describe_table.shape[0]):
        for j in range(df_describe_table.shape[1]):
            table.cell(i+1, j).text = str(df_describe_table.values[i, j])
    
    #nul and non nulls values
    doc.add_heading("Number of Null and Non-Null Values", level=2)
    null_counts = df.isnull().sum()
    non_null_counts = df.notnull().sum()
    counts_df = pd.DataFrame({"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})
    counts_table = doc.add_table(counts_df.shape[0]+1, counts_df.shape[1])
    for i, (col, count) in enumerate(counts_df.items()):
        counts_table.cell(0, i).text = col
        for j, value in enumerate(count):
            counts_table.cell(j+1, i).text = str(value)
            
    # Filtrer les colonnes numériques pour la heatmap
    numeric_columns = df.select_dtypes(include=[float, int]).columns
    numeric_df = df[numeric_columns]
    
    # Ajouter un histogramme pour chaque colonne au rapport
    doc.add_heading("Histograms", level=2)
    for col in numeric_df.columns:
        plt.hist(df[col], bins=20)
        plt.title(col)
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format="png")
        plt.close()
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()

    # Ajouter la heatmap avec seaborn au rapport
    heatmap_df = df[numeric_columns]
    doc.add_heading("Heatmap", level=2)
    plt.figure(figsize=(8, 6))
    sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm')
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    plt.close()
    doc.add_picture(img_buffer, width=Inches(6))
    img_buffer.close()

    # Sauvegarder le rapport en Word
    #doc.save("rapport_analyse.docx")
    # Enregistrer le rapport dans un flux binaire
    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)

    # Générer le lien de téléchargement pour le rapport
    href = f"<a href='data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(report_buffer.read()).decode()}' download='data_analysis_report.docx'>Télécharger le rapport</a>"

    # Afficher le lien de téléchargement dans l'application Streamlit
    st.markdown(href, unsafe_allow_html=True)
    
def main():
    st.set_option('deprecation.showPyplotGlobalUse', False)
    st.title("Exploratory Data analysis of data from a CSV file")

    # Ajout d'un sélecteur de fichiers CSV
    uploaded_file = st.file_uploader("Please choose a CSV file", type=["csv"])

    if uploaded_file is not None:
        # Chargement du fichier CSV dans un DataFrame pandas
        df = pd.read_csv(uploaded_file)
        
        # Affichage du nombre d'enregistrements
        num_records = len(df)
        st.write("Records number :", num_records)
        st.write()
        
        # Affichage du tableau issu de df.describe()
        st.write("General infos:")
        st.write(df.describe())
        st.write()

         # Affichage des 10 premiers lignes
        st.write("First lines:")
        st.write(df.head())
        st.write()   
        
        # Obtenir le nombre de valeurs nulles et non nulles par colonne
        null_counts = df.isnull().sum()
        non_null_counts = df.notnull().sum()

        # Créer un DataFrame à partir des séries null_counts et non_null_counts
        counts_df = pd.DataFrame({"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})

        # Afficher le tableau des nombres de valeurs nulles et non nulles
        st.write("Number of null and non-null values by columns")
        st.table(counts_df)
        
        # Filtrer les colonnes numériques pour la heatmap
        numeric_columns = df.select_dtypes(include=[float, int]).columns
        numeric_df = df[numeric_columns]
        # Affichage d'un histogramme pour chaque colonne
        st.write("Histograms :")
        for col in numeric_df.columns:
            try:
                plt.hist(df[col], bins=20)
                plt.title(col)
                st.pyplot()
                plt.clf()  # Nettoyer la figure après chaque itération
            except:
                print("error")
               # st.write("Cannot draw histogram for column {}".format(col))

        
         # Afficher la heatmap interactive avec Plotly
        st.write("Heatmap :")
        if not numeric_columns.empty:
            fig_heatmap = go.Figure(data=go.Heatmap(z=numeric_df.corr(), x=numeric_df.columns, y=numeric_df.columns, colorscale="RdBu"))
            st.plotly_chart(fig_heatmap)
        else:
            st.write("Aucune colonne numérique disponible pour la heatmap.")
      # Générer et proposer le rapport d'analyse en Word à l'utilisateur
        if st.button("Générer le rapport d'analyse en Word"):
            generate_report(df)
            st.success("Le rapport d'analyse a été généré. Cliquez sur le lien ci-dessous pour le télécharger.")
            #st.markdown(get_file_download_link("rapport_analyse.docx"), unsafe_allow_html=True)
            
def get_file_download_link(file_path):
    """Génère un lien de téléchargement pour le fichier spécifié."""
    with open(file_path, "rb") as file:
        content = file.read()
    href = f"<a href='data:application/octet-stream;base64,{base64.b64encode(content).decode()}' download='{file_path}'>Télécharger le rapport</a>"
    return href


if __name__ == "__main__":
    main()
