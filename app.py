from collections import namedtuple
import altair as alt
import math
import pandas as pd
import seaborn as sns
import streamlit as st
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO
from streamlit.components.v1 import html
from docx import Document
from docx.shared import Inches
import base64
from streamlit.components.v1 import html


def generate_report(df):
    doc = Document()
    doc.add_heading("Analysis report", level=1)

    num_records = len(df)
    doc.add_paragraph(f"Number of records : {num_records}")

    doc.add_heading("General infos:", level=2)
    df_describe_table = df.describe().reset_index()
    df_describe_table.columns = [""] + list(df_describe_table.columns[1:])
    table = doc.add_table(
        df_describe_table.shape[0]+1, df_describe_table.shape[1])
    for i in range(df_describe_table.shape[0]):
        for j in range(df_describe_table.shape[1]):
            table.cell(i+1, j).text = str(df_describe_table.values[i, j])

    doc.add_heading("Number of Null and Non-Null Values", level=2)
    null_counts = df.isnull().sum()
    non_null_counts = df.notnull().sum()
    counts_df = pd.DataFrame(
        {"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})
    counts_table = doc.add_table(counts_df.shape[0]+1, counts_df.shape[1])
    for i, (col, count) in enumerate(counts_df.items()):
        counts_table.cell(0, i).text = col
        for j, value in enumerate(count):
            counts_table.cell(j+1, i).text = str(value)

    numeric_columns = df.select_dtypes(include=[float, int]).columns
    numeric_df = df[numeric_columns]

    doc.add_heading("Histograms", level=2)
    for col in numeric_df.columns:
        plt.hist(df[col], bins=20)
        plt.title(col)
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format="png")
        plt.close()
        doc.add_picture(img_buffer, width=Inches(6))
        img_buffer.close()

    heatmap_df = df[numeric_columns]
    doc.add_heading("Heatmap", level=2)
    plt.figure(figsize=(8, 6))
    sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm')
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    plt.close()
    doc.add_picture(img_buffer, width=Inches(6))
    img_buffer.close()

    report_buffer = BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)

    href = f"<a href='data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(report_buffer.read()).decode()}' download='data_analysis_report.docx'>Download report</a>"

    st.markdown(href, unsafe_allow_html=True)


def main():
    config()
    uploaded_file = st.file_uploader("Please choose a CSV file", type=["csv"])

    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)

        num_records = len(df)
        st.write("Records number :", num_records)
        st.write()

        st.write("General infos")
        st.write(df.describe())
        st.write()

        st.write("First lines")
        st.write(df.head())
        st.write()

        null_counts = df.isnull().sum()
        non_null_counts = df.notnull().sum()

        counts_df = pd.DataFrame(
            {"Number of Null Values": null_counts, "Number of Non-Null Values": non_null_counts})

        st.write("Number of null and non-null values by columns")
        st.table(counts_df)

        numeric_columns = df.select_dtypes(include=[float, int]).columns
        numeric_df = df[numeric_columns]
        for col in numeric_df.columns:
            if col not in (df.index.name, 'index', 'Index'):
                fig = px.histogram(df, x=col, nbins=20, title=col)
                st.plotly_chart(fig)

        st.write("Heatmap")
        if not numeric_columns.empty:
            fig_heatmap = go.Figure(data=go.Heatmap(z=numeric_df.corr(
            ), x=numeric_df.columns, y=numeric_df.columns, colorscale="RdBu"))
            st.plotly_chart(fig_heatmap)
        else:
            st.write("Aucune colonne numérique disponible pour la heatmap.")
        if st.button("Generate analysis report in Word format"):
            generate_report(df)
            st.success(
                "analysis report has been generated, click on link to download.")
        footer_html = """
            <div style="text-align: center; margin-top: 50px;">
                <p>Developped by Ollie </p>
                <p>Date: August 2023</p>
                <p><a href="https://github.com/olivvius/streamlit-example/tree/master">Github link of the app</a></p>
                <p>
                    <a style="color: #fff;" href="https://github.com/olivvius/streamlit-example/tree/master" target="_blank">
                        <img src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://pandas.pydata.org/" target="_blank">
                        <img src="https://upload.wikimedia.org/wikipedia/commons/thumb/e/ed/Pandas_logo.svg/2048px-Pandas_logo.svg.png" alt="Pandas" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://matplotlib.org/" target="_blank">
                        <img src="https://matplotlib.org/stable/_static/logo2_compressed.svg" alt="Matplotlib" height="30" width="30">
                    </a>
                    <a style="color: #fff;" href="https://seaborn.pydata.org/" target="_blank">
                        <img src="https://seaborn.pydata.org/_static/logo-wide-lightbg.svg" alt="Seaborn" height="30" width="120">
                    </a>
                </p>
            </div>
        """
        st.markdown(footer_html, unsafe_allow_html=True)


def config():

    st.set_page_config(
        layout="wide", page_title="Data Analysis", page_icon=":bar_chart:")

    st.markdown(
        """
        <style>
            body {
                background-color: #333;
                color: #fff;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.markdown(
        """
        <style>
            body {
                font-family: 'Helvetica Neue', sans-serif;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.title("Exploratory Data analysis from a CSV file")


if __name__ == "__main__":
    main()
